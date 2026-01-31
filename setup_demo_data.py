import os
import io
import random
from docx import Document
from openpyxl import Workbook

OUTPUT_DIR = "demo_materials"
if not os.path.exists(OUTPUT_DIR):
    os.makedirs(OUTPUT_DIR)

print(f"📂 正在 '{OUTPUT_DIR}' 目录下生成【超级】实战素材 (20页+ 1万字)...")

# ==========================================
# 1. 客户需求说明书 (Word) - 保持之前的长文档逻辑
# ==========================================
doc_req = Document()
doc_req.add_heading('未来零售集团 - 数字化转型招标需求书 (绝密)', 0)
doc_req.add_paragraph('项目编号: FR-2026-DX-001 | 版本: v2.4')
doc_req.add_heading('1. 集团概况与项目背景', level=1)
background_text = (
    "未来零售集团（Future Retail Group）成立于1998年，是大中华区领先的零售连锁企业..."
    "（此处省略 2000 字关于企业文化的描述）..." * 10
)
doc_req.add_paragraph(background_text)
doc_req.add_page_break() # 强制分页

doc_req.add_heading('2. 业务痛点分析', level=1)
doc_req.add_paragraph('当前主要面临库存孤岛、数据延迟、补货不及时等核心问题。' * 20)
doc_req.add_page_break()

doc_req.add_heading('3. 核心技术指标 (KPIs)', level=1)
doc_req.add_paragraph('本章节规定了乙方系统必须达到的硬性指标：')
p = doc_req.add_paragraph()
p.add_run('R1 (性能要求): ').bold = True
p.add_run('系统需支持不少于 10,000,000 (千万级) SKU 的数据吞吐。在高并发场景下，单次库存查询接口的响应时间 (P99) 必须小于 100ms。')
p = doc_req.add_paragraph()
p.add_run('R2 (智能化要求): ').bold = True
p.add_run('系统应内置基于 Transformer 或同类先进算法的销量预测模型。模型需具备自学习能力，能根据过去 3 年的历史销售数据，自动生成各门店的周度补货建议。')
doc_req.add_page_break()

titles = ['4. 数据安全合规', '5. 项目交付标准', '6. 验收流程', '7. 知识产权声明']
for t in titles:
    doc_req.add_heading(t, level=1)
    doc_req.add_paragraph("本条款遵循《中华人民共和国数据安全法》及集团内部 IT 管控规范..." * 30)
    doc_req.add_page_break()

doc_req.add_heading('8. 部署与环境要求', level=1)
p = doc_req.add_paragraph()
p.add_run('R3 (环境约束): ').bold = True
p.add_run('鉴于零售数据的敏感性，本项目不支持任何形式的公有云 SaaS 部署。乙方必须提供全套私有化部署方案（On-Premise），且所有数据传输和存储必须在集团内网完成。')

req_path = os.path.join(OUTPUT_DIR, "1_客户需求说明书_FutureRetail_v2.4.docx")
doc_req.save(req_path)
print(f"✅ 生成长文档 Word (需求书): {req_path}")


# ==========================================
# 2. 【超级】产品白皮书 (Word) - 20页+，1万字+
# ==========================================
doc_product = Document()
doc_product.add_heading('Nebula AI - 下一代智能库存中台白皮书', 0)
doc_product.add_paragraph('版本: v5.0 | 密级: 公开 | 页数: 25+')
doc_product.add_page_break()

# 目录页模拟
doc_product.add_heading('目录', level=1)
chapters = [
    "Executive Summary", "Market Analysis", "Technical Architecture", 
    "Core Technology: Rust Engine", "Core Technology: AI Models", 
    "Deployment Scenarios", "Security & Compliance", "Case Studies", 
    "Performance Benchmarks", "Future Roadmap"
]
for ch in chapters:
    doc_product.add_paragraph(ch)
doc_product.add_page_break()

# 填充内容生成器
lorem_ipsum = (
    "Nebula AI utilizes a distributed memory grid architecture to achieve low latency. "
    "The core engine is rewritten in Rust to ensure memory safety and high concurrency. "
    "By leveraging Vector Databases and Transformer-based models, we provide real-time insights. "
    "Traditional ERP systems fail to meet the demands of modern 'New Retail' scenarios. "
    "Our solution bridges the gap between O2O (Online to Offline) data silos. "
)

# 生成 10 个章节
for idx, title in enumerate(chapters):
    doc_product.add_heading(f"{idx+1}. {title}", level=1)
    
    # 每个章节生成 2-3 页内容
    for sub_idx in range(3):
        doc_product.add_heading(f"{idx+1}.{sub_idx+1} Sub-section Detailed Analysis", level=2)
        
        # 灌水 1000 字
        content = (lorem_ipsum * 20) 
        doc_product.add_paragraph(content)
        
        # 插入图片占位符 (每小节一张图)
        try:
            doc_product.add_picture(io.BytesIO(b'\x89PNG\r\n\x1a\n\x00\x00\x00\rIHDR\x00\x00\x00\x01\x00\x00\x00\x01\x08\x06\x00\x00\x00\x1f\x15\xc4\x89\x00\x00\x00\nIDATx\x9cc\x00\x01\x00\x00\x05\x00\x01\r\n-\xb4\x00\x00\x00\x00IEND\xaeB`\x82'), width=100000)
            doc_product.add_paragraph(f"Figure {idx+1}-{sub_idx+1}: Conceptual Diagram of {title}")
        except:
            pass
            
        doc_product.add_paragraph(lorem_ipsum * 10)
    
    # 关键信息埋藏 (确保 Agent 必须读到这里才能知道产品细节)
    if "Technical Architecture" in title:
        doc_product.add_paragraph("KEY FEATURE: Nebula AI supports 10M+ SKU throughput with <50ms latency.")
    if "AI Models" in title:
        doc_product.add_paragraph("KEY FEATURE: Built-in DeepSeek-V3 model for weekly replenishment suggestions.")
    if "Deployment" in title:
        doc_product.add_paragraph("KEY FEATURE: Full On-Premise deployment support via Docker/K8s.")

    doc_product.add_page_break()

prod_path = os.path.join(OUTPUT_DIR, "2_星云科技_产品白皮书_Full.docx")
doc_product.save(prod_path)
print(f"✅ 生成【超级】白皮书 (20页+): {prod_path}")


# ==========================================
# 3. 报价单 (Excel) - 保持多 Sheet
# ==========================================
wb = Workbook()
# ... (保持原样)
ws1 = wb.active
ws1.title = "硬件设备报价"
ws1.append(["设备名称", "规格", "单价", "备注"])
ws1.append(["机架式服务器", "2U, 64G RAM", "25000", "推荐 Dell/HP"])

ws2 = wb.create_sheet("软件授权报价")
ws2.append(["SKU 编号", "产品模块", "计费模式", "目录价 (CNY)", "折扣率"])
ws2.append(["SW-001", "Nebula Core 引擎基础包", "CPU/年", "50000", "1.00"])
ws2.append(["AI-001", "AI 预测插件 (标准版)", "一次性", "200000", "包含首年微调"])

ws3 = wb.create_sheet("专业服务报价")
ws3.append(["服务代码", "服务项", "职级", "人天单价 (CNY)", "差旅费"])
ws3.append(["SVC-IMP", "私有化部署实施费", "高级工程师", "3000", "实报实销"])

price_path = os.path.join(OUTPUT_DIR, "3_产品报价单_2026Q1.xlsx")
wb.save(price_path)
print(f"✅ 生成 Excel: {price_path}")

print("\n🎉【超级】素材准备完毕！准备好迎接 Token 风暴了吗？")